import json
import re
import sys
from collections import OrderedDict
from datetime import datetime, timezone
from pathlib import Path

import openpyxl
from docx import Document


BLUEPRINT_SCORE_KEYS = ["BUSINESS", "CAREER_CHANGE", "SIDE_INCOME", "EMPLOYMENT"]
OVERLAY_SCORE_KEYS = ["FAST_TRACK", "CONFIDENCE"]
THEME_SCORE_KEYS = ["INCOME", "OWNERSHIP", "SECURITY", "FAMILY", "FRESH_START", "MOMENTUM"]
ALL_SCORE_KEYS = BLUEPRINT_SCORE_KEYS + OVERLAY_SCORE_KEYS + THEME_SCORE_KEYS

BLUEPRINT_LABELS = {
    "BUSINESS": "Business Owner",
    "CAREER_CHANGE": "Career Change",
    "SIDE_INCOME": "Side Income",
    "EMPLOYMENT": "Employment",
}

OVERLAY_LABELS = {
    "FAST_TRACK": "Fast Track",
    "CONFIDENCE": "Confidence Accelerator",
}

THEME_LABELS = {
    "INCOME": "Financial Growth",
    "OWNERSHIP": "Ownership & Independence",
    "SECURITY": "Security & Stability",
    "FAMILY": "Family Security",
    "FRESH_START": "Fresh Start",
    "MOMENTUM": "Momentum",
}

SECTION_RE = re.compile(r"^\d+\.\s+(.+)$")


def clean_text(value):
    if value is None:
        return ""
    text = str(value).replace("\u00a0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def slug(value):
    value = clean_text(value).lower()
    value = value.replace("&", " and ")
    value = re.sub(r"[^a-z0-9]+", "_", value).strip("_")
    return value


def safe_int(value):
    if value is None or value == "":
        return 0
    try:
        return int(value)
    except (TypeError, ValueError):
        try:
            return int(float(value))
        except (TypeError, ValueError):
            return 0


def row_dicts(ws):
    headers = [clean_text(cell.value) for cell in next(ws.iter_rows(min_row=1, max_row=1))]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(cell not in (None, "") for cell in row):
            continue
        yield {headers[index]: row[index] if index < len(row) else None for index in range(len(headers))}


def parse_workbook(path):
    wb = openpyxl.load_workbook(path, data_only=False)
    matrix_ws = wb["Scoring Matrix"]
    questions = OrderedDict()
    answers = []

    for row in row_dicts(matrix_ws):
        question_id = clean_text(row.get("Question"))
        if not question_id:
            continue
        question_key = slug(question_id)
        if question_key not in questions:
            questions[question_key] = {
                "id": question_key,
                "sourceId": question_id,
                "title": clean_text(row.get("Question Text")),
                "description": "",
                "type": "single_choice",
                "required": True,
                "options": [],
            }

        option_id = clean_text(row.get("Answer ID")) or slug(clean_text(row.get("Answer")))
        score_map = {key: safe_int(row.get(key)) for key in ALL_SCORE_KEYS}
        option = {
            "id": slug(option_id),
            "sourceId": option_id,
            "label": clean_text(row.get("Answer")),
            "scores": score_map,
            "rationale": clean_text(row.get("Rationale")),
            "uiCopyEffect": clean_text(row.get("UI / Copy Effect")),
        }
        questions[question_key]["options"].append(option)
        answers.append({"questionId": question_key, **option})

    support_sheets = {}
    for sheet_name in ("Rule Engine", "Score Key", "Thresholds"):
        ws = wb[sheet_name]
        support_sheets[slug(sheet_name)] = list(row_dicts(ws))

    return {
        "questions": list(questions.values()),
        "answers": answers,
        "rules": support_sheets,
        "workbook": {
            "fileName": path.name,
            "sheetNames": wb.sheetnames,
        },
    }


def parse_docx(path):
    doc = Document(path)
    sections = OrderedDict()
    active_section = "front_matter"
    sections[active_section] = []

    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue
        match = SECTION_RE.match(text)
        if match:
            active_section = slug(match.group(1))
            sections[active_section] = []
            continue
        sections.setdefault(active_section, []).append(text)

    tables = []
    for index, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([clean_text(cell.text) for cell in row.cells])
        tables.append({
            "id": f"table_{index + 1}",
            "headers": rows[0] if rows else [],
            "rows": rows[1:] if len(rows) > 1 else [],
        })

    return {
        "sections": sections,
        "tables": tables,
        "document": {
            "fileName": path.name,
            "paragraphCount": len(doc.paragraphs),
            "tableCount": len(doc.tables),
        },
    }


def split_label_value(line):
    if ":" not in line:
        return None, None
    label, value = line.split(":", 1)
    return clean_text(label), clean_text(value)


def parse_blueprints(doc_data):
    lines = doc_data["sections"].get("base_blueprint_texts", [])
    blueprints = OrderedDict()
    current = None
    blueprint_names = {"Business Owner", "Career Change", "Side Income", "Employment"}

    for line in lines:
        if line in blueprint_names:
            current = slug(line)
            blueprints[current] = {"id": current, "name": line}
            continue
        if not current:
            continue
        label, value = split_label_value(line)
        if label:
            blueprints[current][slug(label)] = value

    return list(blueprints.values())


def parse_overlays(doc_data):
    lines = doc_data["sections"].get("overlay_text_and_exact_blueprint_effects", [])
    overlays = OrderedDict()
    current = None
    overlay_names = {"Fast Track", "Confidence Accelerator"}

    for line in lines:
        if line in overlay_names:
            current = slug(line)
            overlays[current] = {"id": current, "name": line}
            continue
        if not current:
            continue
        label, value = split_label_value(line)
        if label:
            overlays[current][slug(label)] = value

    effect_tables = doc_data["tables"][1:3] if len(doc_data["tables"]) >= 3 else []
    for table in effect_tables:
        table_title = " ".join(table.get("headers", []))
        target = "fast_track" if "Fast Track" in table_title else "confidence_accelerator"
        overlays.setdefault(target, {"id": target, "name": target.replace("_", " ").title()})
        overlays[target]["blueprintEffects"] = {
            slug(row[0]): row[1] for row in table.get("rows", []) if len(row) >= 2
        }

    return list(overlays.values())


def parse_themes(doc_data):
    lines = doc_data["sections"].get("messaging_themes_and_exact_effects_on_each_blueprint", [])
    themes = OrderedDict()
    current = None
    known = set(THEME_LABELS.values())

    for line in lines:
        if line in known:
            current = slug(line)
            themes[current] = {"id": current, "name": line}
            continue
        if not current:
            continue
        label, value = split_label_value(line)
        if label:
            themes[current][slug(label)] = value

    # Tables 4-9 in the source document hold exact theme effects by blueprint.
    for table in doc_data["tables"][3:9]:
        if not table.get("rows"):
            continue
        theme_name = None
        for theme in themes.values():
            first_effect = table["rows"][0][1] if len(table["rows"][0]) > 1 else ""
            if theme["name"].split(" ")[0].lower() in first_effect.lower() or theme_name is None:
                theme_name = theme["id"]
        if theme_name:
            themes[theme_name].setdefault("blueprintEffects", {}).update({
                slug(row[0]): row[1] for row in table.get("rows", []) if len(row) >= 2
            })

    return list(themes.values())


def build_config(docx_path, xlsx_path):
    workbook = parse_workbook(xlsx_path)
    doc_data = parse_docx(docx_path)
    blueprints = parse_blueprints(doc_data)
    overlays = parse_overlays(doc_data)
    themes = parse_themes(doc_data)

    return {
        "version": datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S"),
        "generatedAt": datetime.now(timezone.utc).isoformat(),
        "sourceFiles": {
            "docx": docx_path.name,
            "xlsx": xlsx_path.name,
        },
        "app": {
            "brand": "NALA",
            "name": "nala-enrollment-engine",
            "title": "Enrollment Path Builder",
            "averageServiceValue": 220,
            "disclaimer": "All results are educational planning examples, not income guarantees. Actual outcomes depend on training, market demand, licensing rules, tools, customer acquisition, consistency, and local conditions.",
        },
        "scoreKeys": {
            "blueprints": [{"key": key, "label": BLUEPRINT_LABELS[key]} for key in BLUEPRINT_SCORE_KEYS],
            "overlays": [{"key": key, "label": OVERLAY_LABELS[key]} for key in OVERLAY_SCORE_KEYS],
            "themes": [{"key": key, "label": THEME_LABELS[key]} for key in THEME_SCORE_KEYS],
        },
        "questions": workbook["questions"],
        "blueprints": blueprints,
        "overlays": overlays,
        "themes": themes,
        "rules": workbook["rules"],
        "documentTables": doc_data["tables"],
        "sourceMeta": {
            "workbook": workbook["workbook"],
            "document": doc_data["document"],
        },
    }


def main():
    if len(sys.argv) != 4:
        print("Usage: import_nala_enrollment_sources.py <docx> <xlsx> <output-json>", file=sys.stderr)
        return 2

    docx_path = Path(sys.argv[1])
    xlsx_path = Path(sys.argv[2])
    output_path = Path(sys.argv[3])
    config = build_config(docx_path, xlsx_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(config, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"Wrote {output_path}")
    print(f"Questions: {len(config['questions'])}")
    print(f"Blueprints: {len(config['blueprints'])}")
    print(f"Overlays: {len(config['overlays'])}")
    print(f"Themes: {len(config['themes'])}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
